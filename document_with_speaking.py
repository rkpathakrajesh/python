# Create Word Documents With Python

from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

# Picture profile
document = Document()
document.add_picture(r'C:\Users\hp\Desktop\LEARNING-PYTHON\profile_pic.jfif',width=Inches(1.0))

# name phone number and email details
name= input('What is your name? ')
speak('Hello' + name + 'How are you today? ')
phone_number= input('What is your phone number? ')
speak('Your phone number is :' + phone_number + 'Is it correct ? ')
email= input('What is your email? ')
speak('You entered your email id as :' + email + ' I suppose you entered it correctly ')
document.add_paragraph(name + '|'+ phone_number + '|' + email)
# about me
document.add_heading('About Me')
about_me= input('Tell me about yourself ?  ')
speak('You described yourself as ' + about_me + 'Right')
document.add_paragraph(about_me)
# work experience

document.add_heading('Work Experience')
p= document.add_paragraph()
company = input('Enter company : ')
From_date= input('From Date : ')
To_date= input('To Date : ')

# add text to existing paragraph
p.add_run(company + ' ').bold = True
p.add_run(From_date + '-' + To_date + '\n').italic = True
experience_details = input('Describe your experience at ' + company + ' : ')
p.add_run(experience_details)

# more experiences details

while True:
    has_more_experiences = input('Do you have more experiences ? Yes or No : ')
    if has_more_experiences.lower() == 'yes':
        p= document.add_paragraph()
        company = input('Enter company name : ')
        From_date= input('From Date : ')
        To_date= input('To Date : ')

        # add text to existing paragraph
        p.add_run(company + ' ').bold = True
        p.add_run(From_date + '-' + To_date + '\n').italic = True
        experience_details = input('Describe your experience at ' + company + ' : ')
        p.add_run(experience_details)
    else:
        break
# Adding skills to cv

document.add_heading('Add your skills')
p= document.add_paragraph()
skills= input('Enter your skill: ')
time= input('Relevant experience of skill : ')
p.add_run('Name of skill is : ' + skills + '\n').bold= True
p.add_run('Relevant experience of skill ' + skills + ' is :'+ time+': ').bold= True
p.style= 'List Bullet'
while True:
    has_more_skills = input('Has more skills to add? Yes or No : ' )
    if has_more_skills.lower() == 'yes':
        p= document.add_paragraph()
        skills= input('Enter your skill: ')
        time= input('Relevant experience of skill : ')
        p.add_run('Name of skill is : ' + skills + '\n').bold= True
        p.add_run('Relevant experience of skill ' + skills + ' is :'+ time+': ').bold= True
        p.style= 'List Bullet'
    else:
        break
# footer
section = document.sections[0]
footer= section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Amigoscode and Intuit quickbooks course projects'
document.save('cv.docx')







